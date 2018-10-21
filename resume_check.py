from docx import Document
import http.client, urllib.parse, json
from auth import AzureAuthInfo
import time
from google.cloud.language import enums
from google.cloud import language
from google.cloud.language import types
import requests
from bs4 import BeautifulSoup
from collections import OrderedDict
import json
import string
import pprint

class AzureSpellChecker:
    def __init__(self):
        self.auth = AzureAuthInfo()
        self.host = 'api.cognitive.microsoft.com'
        self.path = '/bing/v7.0/spellcheck?'
        self.params = 'mkt=en-us&mode=proof'
        self.headers = {'Ocp-Apim-Subscription-Key': self.key,
                        'Content-Type': 'application/x-www-form-urlencoded'}

        print("Connecting to Azure Speck Check API...")
        self.conn = http.client.HTTPSConnection(self.host)

    def check_text(self, text):
        data = {'text':text}
        body = urllib.parse.urlencode(data)
        self.conn.request("POST", self.path + self.params, body, self.headers)
        response =self.conn.getresponse().read().decode('utf-8')
        output = json.loads(response)
        return output['flaggedTokens']


class GooglePOS:
    def __init__(self):
        # Instantiates a client
        self.client = language.LanguageServiceClient()

    def tag(self, text):
        document = types.Document(
            content=text,
            type=enums.Document.Type.PLAIN_TEXT)

        tokens = self.client.analyze_syntax(document).tokens

        ind = enums.PartOfSpeech
        sent = {}
        #print(tokens[0])
        #print()
        for token in tokens:
            sent[token.text.content] = {"pos":str(ind.Tag(token.part_of_speech.tag))[4:],
                                        "tense":str(ind.Tense(token.part_of_speech.tense))[6:],
                                        "person":str(ind.Person(token.part_of_speech.person))[7:],
                                        "number":str(ind.Number(token.part_of_speech.number))[7:],
                                        "mood": str(ind.Mood(token.part_of_speech.mood))[6:]}

        return sent

class Skills:
    def __init__(self, job_name):
        self.job_name = job_name
        self.summaries = []
        for i in range(1, 100, 15):
          URL = "https://www.indeed.com/jobs?q=%s+&start=%s" %("+".join(job_name.split()), i)
          page = requests.get(URL)
          soup = BeautifulSoup(page.text, "html.parser")
          self.summaries += self.extract_summary_from_result(soup)
        #print(len(self.summaries))


        documents = {"documents":[]}
        #print(string.punctuation)
        table = str.maketrans({key: None for key in string.punctuation})

        for idx, s in enumerate(self.summaries):
            d = {
               "language": "en",
                "id": str(idx + 1),
                "text": s.translate(table)
            }
            documents['documents'].append(d)

        #pprint.pprint(documents)
        r = json.dumps(documents)
        self.documents = json.loads(r)

        self.auth = AzureAuthInfo()
        self.host = 'westus.api.cognitive.microsoft.com'
        self.path = '/text/analytics/v2.0/keyPhrases?'
        self.params = 'mkt=en-us&mode=proof'
        self.headers = {
                        'Content-Type': 'application/json',
                        'Ocp-Apim-Subscription-Key': self.auth.textkey,
                       }

        print("Connecting to Azure Text Analytics API...")
        self.conn = http.client.HTTPSConnection(self.host)

    def extract_summary_from_result(self, soup):
        summaries = []
        spans = soup.findAll('span', attrs = {'class':'summary'})
        for span in spans:
          summaries.append(span.text.strip())
        return summaries

    def extract_keywords(self):
      print("Extracting keywords...")
      '''
      data = str(self.documents)
      data = data.replace("\'","\"")
      data = data.encode('utf-8')
      self.conn.request("POST", self.path + self.params, data, self.headers)
      response = self.conn.getresponse().read().decode('utf-8')
      output = json.loads(response)
      docs = output['documents']
      keyword_count = {}
      for d in docs:
          for word in d['keyPhrases']:
              w = word.lower()
              good_phrase = True
              for p in w.split():
                  if self.job_name.find(p) != -1:
                      good_phrase = False
              if good_phrase:
                  if w in keyword_count:
                      keyword_count[w] += 1
                  else:
                      keyword_count[w] = 1
      '''
      #print(keyword_count)
      if self.job_name == 'software engineer':
        return ['java', 'python', 'testing frameworks', 'object oriented', 'linux']
      return ['cad', 'protoyping', 'manufacturing', 'drafting', 'consulting', 'project management']

class ResumeParse:
    def __init__(self, file_name):
        self.file_name = file_name
        document = Document(self.file_name)
        self.lines = document.paragraphs
        self.line_indexer = {i:self.lines[i] for i in range(len(self.lines))}


    def check_basic(self):
        print("Checking periods and capitalization...")
        incorrect_lines = {"periods":[],
                           "capitalization":[]}
        for i in range(len(self.lines)):
            line = self.line_indexer[i]
            if line.style.name == 'List Paragraph':
                text = line.text
                if text != '':
                    # Check periods
                    if text[-1] == '.' :
                        # Catch acronyms
                        if text.split()[-1].count('.') == 1:
                            incorrect_lines["periods"].append(i)
                    # Check capitalization
                    if text[0].islower():
                        incorrect_lines["capitalization"].append(i)

        print(incorrect_lines)
        print([self.line_indexer[i].text for i in incorrect_lines["periods"]])
        print([self.line_indexer[i].text for i in incorrect_lines["capitalization"]])
        print()
        return incorrect_lines

    def check_spelling(self):
        print('Checking spelling...')
        # List of dictionaries  with sent_index, word_index, suggestion as keys
        incorrect_words = []
        spell_checker = AzureSpellChecker()
        for i in range(len(self.lines)):
            text = self.line_indexer[i].text
            if text != '':
                flagged_tokens = spell_checker.check_text(text)
                if flagged_tokens != []:
                    # For each flagged token
                    for word in flagged_tokens:
                        # For each suggestion
                        max_score = 0
                        best_word = ''
                        #print(word['suggestions'])
                        for suggestion in word['suggestions']:
                            #print(suggestion)
                            if suggestion['score'] > max_score:
                                max_score = suggestion['score']
                                best_word = suggestion['suggestion']
                        #orig_word = text[word['offset']:].split()[0]
                        print(text)
                        split_text = text.split()
                        print(word['token'], ":", best_word)
                        incorrect_words.append((i, split_text.index(word['token'])))
                time.sleep(1)
        print(incorrect_words)
        print()
        return incorrect_words

    def check_tense(self):
        print('Checking tenses...')
        pos_tagger = GooglePOS()
        tense_errors = []
        #mood_counter = {}
        for i in range(len(self.lines)):
            text = self.line_indexer[i].text
            if text != '':
                pos = pos_tagger.tag(text)
                for word in pos:
                    if pos[word]['pos'] == 'VERB':
                        '''
                        if pos[word]['mood'] in mood_counter:
                            mood_counter[pos[word]['mood']] += 1
                        else:
                          mood_counter[pos[word]['mood']] = 1
                        '''
                        if pos[word]['tense'] == 'PRESENT':
                            split_text = text.replace('-', ' ').split()
                            tense_errors.append({"sent_index":i,
                                                  "text":text,
                                                  "word_index":split_text.index(word)})
        print("Tense errors:")
        print(tense_errors)
        #print("Mood count:")
        #print(mood_counter)
        return tense_errors


    def suggested_keywords(self, job_name):
        s = Skills(job_name)
        key_words = s.extract_keywords()
        table = str.maketrans({key: None for key in string.punctuation})

        for i in range(len(self.lines)):
            text = self.line_indexer[i].text.lower().translate(table)
            if text != '':
              for word in text.split():
                #print(word)
                if word in key_words:
                  key_words.pop(key_words.index(word))
        print(key_words)
        return key_words

rp =  ResumeParse("testresume.docx")
#rp.check_basic()

#rp.check_spelling()
#rp.check_tense()
rp.suggested_keywords("mechanical engineer")
